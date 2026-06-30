#!/usr/bin/env python3
"""
生成汽车销售信息管理系统课程设计报告 Word 文档
参照课设模板格式：宋体正文、黑体标题、A4页面
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
#  全局样式设置
# ============================================================

# 设置默认字体（正文：宋体小四 12pt）
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
# 设置中文字体
rpr = style.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn('w:eastAsia'), '宋体')

# 页面设置 A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


# ============================================================
#  辅助函数
# ============================================================

def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=12, bold=False, color=None):
    """设置 run 的中英文字体"""
    run.font.name = en_font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), cn_font)


def add_heading_custom(text, level=1):
    """添加自定义标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:  # 大标题
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Arial', 22, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:  # 一级标题
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Arial', 16, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:  # 二级标题
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Arial', 14, bold=True)
    elif level == 3:  # 三级标题
        run = p.add_run(text)
        set_run_font(run, '黑体', 'Arial', 12, bold=True)
    return p


def add_body(text, indent=True, bold=False, align='left'):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)  # 1.5倍行距约22pt
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', 12, bold=bold)
    return p


def add_code(text):
    """添加代码块（等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, 'Courier New', 'Courier New', 10)
    # 灰色背景
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._element.get_or_add_pPr().append(shading)
    return p


def add_page_break():
    doc.add_page_break()


# ============================================================
#  封面页
# ============================================================

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('课  程  设  计')
set_run_font(run, '黑体', 'Arial', 26, bold=True)

doc.add_paragraph()
doc.add_paragraph()

# 封面信息表
info_items = [
    ('课程名称：', 'C语言程序设计课程设计'),
    ('设计课题：', '汽车销售信息管理系统设计'),
    ('指导教师：', ''),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(label + value)
    set_run_font(run, '宋体', 'Times New Roman', 16)

for _ in range(4):
    doc.add_paragraph()

for label, value in [('专　业：', '　　　　　　　　'), ('班　级：', '　　　　　　　　')]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(label + value)
    set_run_font(run, '宋体', 'Times New Roman', 14)

for label, value in [('姓　名：', '　　　　　　　　'), ('学　号：', '　　　　　　　　')]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(label + value)
    set_run_font(run, '宋体', 'Times New Roman', 14)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('二〇二六年 6 月    日')
set_run_font(run, '宋体', 'Times New Roman', 14)

add_page_break()


# ============================================================
#  课程设计任务书
# ============================================================

add_heading_custom('课程设计任务书', 0)
doc.add_paragraph()

add_heading_custom('一、课程设计目的', 1)
purposes = [
    '1、巩固并加深学生对C语言程序设计知识的理解；',
    '2、培养学生使用链表等数据结构进行程序设计的能力；',
    '3、进一步掌握模块化编程思想和函数封装方法；',
    '4、提高运用C语言解决实际问题的能力；',
    '5、初步掌握开发小型实用软件的基本方法，能独立设计、实现具有实际功能的小系统；',
    '6、掌握书写程序设计开发文档的能力（书写课程设计实验报告）。',
]
for t in purposes:
    add_body(t)

add_heading_custom('二、课程设计要求', 1)
reqs = [
    '1、根据分组情况对规定的课程设计题目要求进行分析、设计；',
    '2、利用C语言独立编写程序代码，并调试程序使其能正确运行；',
    '3、源程序要有适当的注释，使程序容易阅读；',
    '4、设计完成的软件要便于操作和使用，采用模块化设计方法；',
    '5、要求系统的用户界面尽量简单、合理，鼓励学生自动增加新功能模块（视情况可另行加分）；',
    '6、完成并提交课程设计报告，具体要求见相关说明文档。',
]
for t in reqs:
    add_body(t)

add_heading_custom('三、课程设计内容', 1)
add_body('汽车销售信息管理系统设计', bold=True)

add_heading_custom('1、问题描述', 2)
add_body('汽车销售信息管理系统用于管理汽车4S店或经销商的日常业务数据。系统需要管理的实体包括：轿车信息（编号、型号、颜色、生产厂家、出厂日期、价格）、员工信息（编号、姓名、性别、年龄、籍贯、学历）、客户信息（编号、名称、联系方式、地址、业务联系记录）、销售信息（记录编号、销售日期、汽车编号、型号、颜色、数量、经手人、客户、总价）。销售记录关联轿车、员工、客户三个实体，需要实现跨表关联查询和自动填充。')

add_heading_custom('2、功能要求', 2)
funcs = [
    '（1）轿车信息管理：添加、删除、修改、查询（按编号/型号/价格范围）、显示全部',
    '（2）员工信息管理：添加、删除、修改、查询（按编号/姓名/学历）、显示全部',
    '（3）客户信息管理：添加、删除、修改、查询（按编号/名称）、显示全部、追加业务联系记录',
    '（4）销售信息管理：添加（自动关联车辆/员工/客户信息）、删除、查询（按编号/型号/经手人/日期范围）、显示全部',
    '（5）销售报表统计：按型号统计销售数量和金额、按经手人统计销售数量和金额',
    '（6）数据持久化：退出系统时自动保存全部数据到文件，启动时自动加载',
]
for t in funcs:
    add_body(t)

add_heading_custom('3、问题的解决方案', 2)
steps = [
    '（1）应用系统分析，建立该系统的功能模块框图以及界面的组织和设计；',
    '（2）分析系统中的各个实体及它们之间的关系；',
    '（3）根据问题描述，设计各实体的结构体（struct）和链表存储方案；',
    '（4）完成各模块中增删改查函数的定义；',
    '（5）完成文件读写功能的实现（二进制存储）；',
    '（6）完成主菜单和各子菜单的交互逻辑；',
    '（7）功能调试；',
    '（8）完成系统总结报告。',
]
for t in steps:
    add_body(t)

add_page_break()


# ============================================================
#  前言
# ============================================================

add_heading_custom('前  言', 0)
doc.add_paragraph()

add_body('本汽车销售信息管理系统基于C语言开发，采用单向链表作为核心数据结构，二进制文件作为数据持久化方式。系统针对汽车销售行业日常业务管理需求，围绕轿车、员工、客户、销售四类核心数据的增删改查场景，实现完整的信息管理功能。')

add_body('项目采用模块化设计，将系统划分为轿车管理、员工管理、客户管理、销售管理、数据I/O和通用工具六个独立模块，通过头文件声明接口、源文件实现逻辑的方式组织代码。数据层使用链表存储，文件层使用 fwrite/fread 二进制读写，形成内存链表＋磁盘文件的双重存储机制。开发过程紧扣C语言指针操作、链表遍历、结构体封装、文件I/O等核心知识点，把课本理论转化为可独立运行的实用软件。')

add_body('系统在设计阶段完善了输入校验逻辑，对编号唯一性、空输入等边界情况做了处理；销售记录添加时自动关联车辆、员工、客户信息，减少重复录入；销售报表支持按型号和按经手人两个维度统计，满足基本的业务分析需求。本次课程设计完整走完需求分析、数据结构设计、模块开发与调试测试全流程，也为后续更复杂的数据管理系统开发积累了实践经验。')

add_page_break()


# ============================================================
#  目录
# ============================================================

add_heading_custom('目  录', 0)
doc.add_paragraph()

toc_items = [
    ('一、需求分析', 1),
    ('    1.1 项目背景 ............................................................ 1', 2),
    ('    1.2 功能需求 ............................................................ 1', 2),
    ('    1.3 数据需求 ............................................................ 1', 2),
    ('    1.4 运行环境 ............................................................ 2', 2),
    ('二、概要设计', 1),
    ('    2.1 系统总体架构框图 .................................................. 3', 2),
    ('    2.2 数据结构关系图 .................................................... 4', 2),
    ('    2.3 主程序业务流程图 .................................................. 5', 2),
    ('三、详细设计', 1),
    ('    3.1 系统功能模块划分 .................................................. 6', 2),
    ('    3.2 数据结构详细设计 .................................................. 8', 2),
    ('    3.3 核心函数功能与参数说明 ............................................ 9', 2),
    ('    3.4 模块间接口关系 ................................................... 11', 2),
    ('四、源程序清单和执行结果', 1),
    ('    4.1 项目文件结构 ..................................................... 12', 2),
    ('    4.2 各模块核心源码 ................................................... 13', 2),
    ('    4.3 各功能运行结果说明 ............................................... 20', 2),
    ('五、用户使用说明', 1),
    ('    5.1 项目编译与运行步骤 ............................................... 21', 2),
    ('    5.2 系统各功能使用说明 ............................................... 21', 2),
    ('    5.3 软件使用注意事项 ................................................. 22', 2),
    ('六、调试与测试', 1),
    ('    6.1 功能测试用例 ..................................................... 23', 2),
    ('    6.2 常见bug及解决办法 ............................................... 24', 2),
    ('七、总结、附录和参考资料', 1),
    ('    7.1 项目开发总结 ..................................................... 25', 2),
    ('    7.2 附录 ............................................................. 26', 2),
    ('    7.3 参考文献 ......................................................... 27', 2),
]

for text, level in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, '黑体', 'Arial', 12, bold=True)
    else:
        set_run_font(run, '宋体', 'Times New Roman', 12)

add_page_break()


# ============================================================
#  一、需求分析
# ============================================================

add_heading_custom('一、需求分析', 0)
doc.add_paragraph()

add_heading_custom('1.1 项目背景', 1)
add_body('传统汽车销售业务中，车辆库存、销售人员、客户信息和交易记录往往分散在纸质档案或简易电子表格中，查询效率低、数据易丢失、统计工作繁琐。为实现电子化管理，采用C语言开发本系统，集中管理轿车、员工、客户和销售四类数据，完成信息增删改查与报表统计，提升业务管理效率。')

add_heading_custom('1.2 功能需求', 1)
features = [
    '1. 轿车管理：添加车辆（编号唯一性检查）、删除（确认后删除）、修改（回车保留原值）、查询（支持按编号精确查询、按型号模糊查询、按价格范围查询）、显示全部车辆列表。',
    '2. 员工管理：添加员工（编号查重）、删除、修改、查询（支持按编号、按姓名模糊、按学历筛选）、显示全部。',
    '3. 客户管理：添加客户（编号自动生成 C0001/C0002...）、删除、修改、查询（支持按编号、按名称模糊、查看含联系记录的完整信息）、显示全部、追加业务联系记录（自动添加日期戳）。',
    '4. 销售管理：添加销售记录（编号自动生成 S0001/S0002...，自动关联车辆/员工/客户并填充冗余字段，自动计算总价）、删除、查询（支持按编号、按车型、按经手人、按日期范围）、显示全部、生成销售报表（按型号统计＋按经手人统计，同时导出文本文件）。',
    '5. 数据持久化：系统启动时自动从文件加载数据，退出时自动保存全部数据到二进制文件。',
]
for t in features:
    add_body(t)

add_heading_custom('1.3 数据需求', 1)

add_heading_custom('1. 轿车结构体（Car）', 3)
add_body('id: 编号（char[64]，主键唯一）')
add_body('model: 型号（char[64]）')
add_body('color: 颜色（char[64]）')
add_body('manufacturer: 生产厂家（char[64]）')
add_body('date: 出厂日期（char[64]，格式 YYYY-MM-DD）')
add_body('price: 价格（double）')
add_body('next: 链表指针')

add_heading_custom('2. 员工结构体（Employee）', 3)
add_body('id: 编号（char[64]，主键唯一）')
add_body('name: 姓名（char[64]）')
add_body('gender: 性别（char[4]）')
add_body('age: 年龄（int）')
add_body('hometown: 籍贯（char[64]）')
add_body('education: 学历（enum: 1-高中 2-大专 3-本科 4-硕士 5-博士）')
add_body('next: 链表指针')

add_heading_custom('3. 客户结构体（Customer）', 3)
add_body('id: 编号（char[64]，自动生成 C0001...）')
add_body('name: 客户名称（char[64]）')
add_body('phone: 联系方式（char[64]）')
add_body('address: 地址（char[256]）')
add_body('contact_log: 业务联系记录（char[256]，格式 [YYYY-MM-DD] 内容; 追加）')
add_body('next: 链表指针')

add_heading_custom('4. 销售结构体（Sale）', 3)
add_body('id: 记录编号（char[64]，自动生成 S0001...）')
add_body('date: 销售日期（char[64]，格式 YYYY-MM-DD）')
add_body('car_id: 关联车辆编号（char[64]）')
add_body('car_model: 型号（char[64]，冗余存储）')
add_body('color: 颜色（char[64]，冗余存储）')
add_body('quantity: 数量（int）')
add_body('emp_id: 经手人编号（char[64]）')
add_body('emp_name: 经手人姓名（char[64]，冗余存储）')
add_body('cust_id: 客户编号（char[64]）')
add_body('cust_name: 客户名称（char[64]，冗余存储）')
add_body('total_price: 总价（double，= 单价 × 数量）')
add_body('next: 链表指针')

add_heading_custom('5. 实体间关系', 3)
add_body('Sale.car_id → Car.id（一对多，一辆车可被多次销售）')
add_body('Sale.emp_id → Employee.id（一对多，一个员工可经手多条记录）')
add_body('Sale.cust_id → Customer.id（一对多，一个客户可有多个订单）')

add_heading_custom('1.4 运行环境', 1)
add_body('开发语言：C11')
add_body('编译工具链：GCC + CMake')
add_body('操作系统：Linux（Arch Linux），兼容 Windows')
add_body('构建系统：CMake 3.10+')
add_body('运行方式：命令行终端')

add_page_break()


# ============================================================
#  二、概要设计
# ============================================================

add_heading_custom('二、概要设计', 0)
doc.add_paragraph()

add_heading_custom('2.1 系统总体架构框图', 1)
add_body('系统采用分层模块化设计，分为三层：')
add_body('用户交互层（menu.c / main.c）：主菜单、子菜单，接收用户输入并调度功能')
add_body('业务模块层（car / employee / customer / sale）：各自管理一类数据的增删改查')
add_body('基础支撑层（data_io / utils）：统一文件读写入口、公共工具函数')
doc.add_paragraph()
add_code('''                    ┌─────────────────────────┐
                    │       用户交互层         │
                    │    (menu.c / main.c)     │
                    └──────────┬──────────────┘
                               │
              ┌────────────────┼────────────────┐
              │                │                │
    ┌─────────▼────────┐ ┌────▼─────┐ ┌───────▼──────┐
    │   轿车管理模块    │ │ 员工管理 │ │  客户管理     │
    │    (car.c/h)     │ │(employee)│ │ (customer)    │
    └─────────┬────────┘ └────┬─────┘ └───────┬──────┘
              │                │                │
    ┌─────────▼────────┐ ┌────▼─────┐ ┌───────▼──────┐
    │   销售管理模块    │ │ data_io  │ │  utils       │
    │   (sale.c/h)     │ │ (data_io)│ │  (utils)     │
    └──────────────────┘ └──────────┘ └──────────────┘''')
add_body('（图2-1 系统总体架构框图）', align='center')

add_heading_custom('2.2 数据结构关系图', 1)
add_body('四个链表独立存储，销售记录通过编号引用关联实体：')
doc.add_paragraph()
add_code('''  ┌──────────┐    ┌──────────┐    ┌───────────┐
  │  Car     │    │ Employee │    │ Customer  │
  │ 链表     │    │ 链表     │    │ 链表      │
  │ car_head │    │ emp_head │    │ cust_head │
  └────┬─────┘    └────┬─────┘    └─────┬─────┘
       │car_id         │emp_id          │cust_id
       └───────────────┼────────────────┘
                       │
              ┌────────▼────────┐
              │    Sale         │
              │    链表         │
              │  sale_head      │
              └─────────────────┘''')
add_body('（图2-2 数据结构关系图）', align='center')

add_heading_custom('2.3 主程序业务流程图', 1)
doc.add_paragraph()
add_code('''程序启动
  │
  ▼
load_all_data() ──→ car_load() / emp_load() / cust_load() / sale_load()
  │
  ▼
main_menu() 循环：
  │
  ├── 1 → car_menu()      [添加/删除/修改/查询/显示]
  ├── 2 → employee_menu() [添加/删除/修改/查询/显示]
  ├── 3 → customer_menu() [添加/删除/修改/查询/显示/追加记录]
  ├── 4 → sale_menu()     [添加/删除/查询/显示/报表]
  ├── 5 → save_all_data()
  └── 0 → save_all_data() → 退出''')
add_body('（图2-3 主程序业务流程图）', align='center')

add_page_break()


# ============================================================
#  三、详细设计
# ============================================================

add_heading_custom('三、详细设计', 0)
doc.add_paragraph()

add_heading_custom('3.1 系统功能模块划分', 1)

add_heading_custom('1. 主控模块（main.c + menu.c）', 3)
add_body('main.c：程序入口，调用 load_all_data() 加载数据后进入 main_menu() 循环。')
add_body('menu.c：主菜单显示与分发，包含6个选项（轿车/员工/客户/销售/保存/退出）。退出时自动保存数据。')

add_heading_custom('2. 轿车管理模块（car.c/h）', 3)
add_body('核心数据结构：Car 单向链表（含尾指针 car_tail 优化尾部追加）')
add_body('功能函数：car_add()（添加，编号查重，尾插法）、car_delete()（删除，维护头尾指针）、car_modify()（修改，回车保留原值）、car_query()（子菜单，按编号/型号模糊/价格范围）、car_list_all()、car_find()、car_save/load()。')

add_heading_custom('3. 员工管理模块（employee.c/h）', 3)
add_body('核心数据结构：Employee 单向链表 + Education 枚举')
add_body('功能函数：emp_add()（编号查重，学历枚举验证）、emp_delete()、emp_modify()、emp_query()（按编号/姓名模糊/学历筛选）、emp_list_all()、emp_find()、emp_save/load()、edu_to_str()。')

add_heading_custom('4. 客户管理模块（customer.c/h）', 3)
add_body('核心数据结构：Customer 单向链表 + 自增编号计数器')
add_body('功能函数：cust_add()（编号自动生成 C0001...）、cust_delete()、cust_modify()、cust_query()（按编号/名称模糊/查看完整信息）、cust_list_all()、cust_add_contact_log()（追加联系记录，自动日期戳）、cust_find()、cust_save/load()（加载时同步更新编号计数器）。')

add_heading_custom('5. 销售管理模块（sale.c/h）', 3)
add_body('核心数据结构：Sale 单向链表 + SaleStat 统计结构体')
add_body('功能函数：sale_add()（自动关联车辆/员工/客户，填充冗余字段，计算总价）、sale_delete()、sale_query()（按编号/车型/经手人/日期范围）、sale_list_all()、sale_report()（按型号＋按经手人统计，导出 report.txt）、sale_find()、sale_save/load()。')

add_heading_custom('6. 数据I/O模块（data_io.c/h）', 3)
add_body('load_all_data()：启动时调用，按序加载四类数据')
add_body('save_all_data()：退出/手动保存时调用，按序保存四类数据')
add_body('ensure_data_dir()：确保 data/ 目录存在（跨平台 mkdir）')

add_heading_custom('7. 工具模块（utils.c/h）', 3)
add_body('clear_screen()：清屏（Linux clear / Windows cls）')
add_body('pause_screen()：按回车继续')
add_body('safe_gets()：安全读取一行（处理缓冲区残留）')
add_body('read_int()：安全读取整数')
add_body('get_current_date()：获取当前日期字符串（YYYY-MM-DD）')

add_heading_custom('3.2 数据结构详细设计', 1)

add_body('1. 链表采用尾插法（而非头插法）：每个模块维护 head 和 tail 两个指针，新节点追加到尾部。好处是列表中数据的显示顺序与添加顺序一致，用户体验更好。')

add_body('2. 删除节点时同时维护 tail 指针：当删除的是尾节点时，将 tail 指向前驱节点，保证 tail 始终有效。')

add_body('3. 销售记录的冗余存储：car_model、color、emp_name、cust_name 在添加销售记录时从关联表查出后复制保存。这样做的好处是报表查询和显示时不依赖其他模块的数据完整性，即使车辆或员工被删除，历史销售记录仍然完整可读。')

add_body('4. 客户编号和销售编号自动生成：使用静态计数器 cust_id_counter 和 sale_id_counter，每次自增后格式化输出（snprintf）。加载数据时扫描已有编号更新计数器最大值，避免重启后编号冲突。')

add_heading_custom('3.3 核心函数功能与参数说明', 1)

add_heading_custom('1. 通用查找函数模式', 3)
add_body('Car *car_find(const char *id)：遍历 car_head 链表，strcmp 比较 id，找到返回指针，否则返回 NULL')
add_body('Employee *emp_find(const char *id)：同上')
add_body('Customer *cust_find(const char *id)：同上')
add_body('Sale *sale_find(const char *id)：同上')

add_heading_custom('2. 添加函数模式', 3)
add_body('car_add()：分配节点→输入编号→查重→输入其余字段→尾插→提示成功')
add_body('emp_add()：类似，额外处理学历枚举输入验证')
add_body('cust_add()：编号自动生成→跳过已存在编号→输入其余字段→追加首次联系记录→尾插')
add_body('sale_add()：编号自动生成→输入日期（默认当日）→关联车辆（自动填充型号/颜色/单价）→输入数量→关联员工→关联客户→计算总价→尾插')

add_heading_custom('3. 删除函数模式', 3)
add_body('所有删除函数遵循统一流程：输入编号→遍历查找（记录 prev）→显示记录→确认→摘链（分 prev==NULL 和 prev!=NULL 两种情况）→更新 tail→free→提示')

add_heading_custom('4. 修改函数模式', 3)
add_body('所有修改函数遵循统一流程：输入编号→查找→显示当前值→逐字段提示"回车保留原值"→仅当输入非空时才更新')

add_heading_custom('5. 查询函数', 3)
add_body('各模块查询均为子菜单循环，支持多种查询方式：')
add_body('car_query：编号精确 / 型号模糊(strstr) / 价格范围')
add_body('emp_query：编号精确 / 姓名模糊(strstr) / 学历精确')
add_body('cust_query：编号精确 / 名称模糊(strstr) / 查看完整信息')
add_body('sale_query：编号精确 / 车型模糊(strstr) / 经手人模糊(strstr) / 日期范围(strcmp)')

add_heading_custom('6. 销售报表函数', 3)
add_body('sale_report()：遍历 sale_head 链表；用 SaleStat[100] 数组分别按型号和按经手人做去重汇总（sale_add_stat）；打印两个统计表到 stdout；同时将报表写入 data/report.txt。')

add_heading_custom('3.4 模块间接口关系', 1)
add_body('模块间依赖关系：')
add_body('menu.c → car.h, employee.h, customer.h, sale.h, data_io.h, utils.h')
add_body('data_io.c → car.h, customer.h, employee.h, sale.h')
add_body('sale.c → car.h, employee.h, customer.h（关联查询需要）')
add_body('car.c, employee.c, customer.c → utils.h')
add_body('main.c → menu.h, data_io.h')
doc.add_paragraph()
add_body('跨模块调用：')
add_body('sale_add() 调用 car_find()、emp_find()、cust_find() 实现自动关联')
add_body('menu.c 通过各模块的 _menu() 函数进入子菜单')
add_body('data_io.c 通过各模块的 _save()/_load() 函数实现统一读写')

add_page_break()


# ============================================================
#  四、源程序清单和执行结果
# ============================================================

add_heading_custom('四、源程序清单和执行结果', 0)
doc.add_paragraph()

add_heading_custom('4.1 项目文件结构', 1)
add_code('''Automobile_Sales_Information_Management_System/
├── CMakeLists.txt          # CMake构建配置
├── include/                # 头文件目录
│   ├── car.h               # 轿车模块接口
│   ├── customer.h          # 客户模块接口
│   ├── employee.h          # 员工模块接口
│   ├── sale.h              # 销售模块接口
│   ├── data_io.h           # 数据I/O模块接口
│   ├── utils.h             # 工具函数接口
│   └── menu.h              # 菜单接口
├── src/                    # 源文件目录
│   ├── main.c              # 程序入口
│   ├── menu.c              # 主菜单
│   ├── car.c               # 轿车管理实现
│   ├── employee.c          # 员工管理实现
│   ├── customer.c          # 客户管理实现
│   ├── sale.c              # 销售管理实现
│   ├── data_io.c           # 数据I/O实现
│   └── utils.c             # 工具函数实现
└── data/                   # 数据文件目录（运行时生成）
    ├── cars.dat            # 车辆数据（二进制）
    ├── employees.dat       # 员工数据（二进制）
    ├── customers.dat       # 客户数据（二进制）
    ├── sales.dat           # 销售数据（二进制）
    └── report.txt          # 销售报表（文本）''')

add_heading_custom('4.2 各模块核心源码', 1)
add_body('（此处粘贴各源文件的完整源码。以下展示关键代码片段，完整源码详见仓库 src/ 和 include/ 目录。）')

add_heading_custom('(1) 链表尾插法（car.c）', 3)
add_code('''if (car_head == NULL) {
    car_head = new_node;
    car_tail = new_node;
} else {
    car_tail->next = new_node;
    car_tail = new_node;
}''')

add_heading_custom('(2) 安全输入函数（utils.c）', 3)
add_code('''void safe_gets(char *buf, int size) {
    if (fgets(buf, size, stdin) == NULL) {
        buf[0] = '\\0';
        return;
    }
    size_t len = strlen(buf);
    if (len > 0 && buf[len - 1] == '\\n') {
        buf[len - 1] = '\\0';
    } else {
        int ch;
        while ((ch = getchar()) != '\\n' && ch != EOF) {}
    }
}''')

add_heading_custom('(3) 销售记录自动关联（sale.c）', 3)
add_code('''Car *car = car_find(node->car_id);
if (car != NULL) {
    strcpy(node->car_model, car->model);
    strcpy(node->color, car->color);
    unit_price = car->price;
    printf("已关联车辆: %s / %s / 单价 %.2f\\n",
           node->car_model, node->color, unit_price);
} else {
    printf("未找到车辆信息，请手动输入型号和颜色。\\n");
    printf("汽车型号: ");
    safe_gets(node->car_model, sizeof(node->car_model));
    printf("颜色: ");
    safe_gets(node->color, sizeof(node->color));
    unit_price = read_nonnegative_double("单价: ");
}''')

add_heading_custom('(4) 销售报表统计（sale.c）', 3)
add_code('''static void sale_add_stat(SaleStat stats[], int *count, const char *name,
                          int qty, double amount) {
    const char *key = name[0] ? name : "未填写";
    for (int i = 0; i < *count; i++) {
        if (strcmp(stats[i].name, key) == 0) {
            stats[i].qty += qty;
            stats[i].amount += amount;
            return;
        }
    }
    if (*count >= 100) return;
    snprintf(stats[*count].name, sizeof(stats[*count].name), "%s", key);
    stats[*count].qty = qty;
    stats[*count].amount = amount;
    (*count)++;
}''')

add_heading_custom('(5) 二进制文件读写（data_io.c）', 3)
add_code('''void load_all_data(void) {
    ensure_data_dir();
    car_load();
    emp_load();
    cust_load();
    sale_load();
}

void save_all_data(void) {
    ensure_data_dir();
    car_save();
    emp_save();
    cust_save();
    sale_save();
}''')

add_heading_custom('4.3 各功能运行结果说明', 1)
add_body('（此处添加实际运行截图：主菜单界面、各子菜单界面、添加记录、查询结果、销售报表等。截图应覆盖以下场景：）')
add_body('1. 主菜单启动界面')
add_body('2. 车辆添加成功提示')
add_body('3. 重复编号添加拒绝提示')
add_body('4. 车辆查询结果（按型号模糊查询）')
add_body('5. 员工按学历筛选结果')
add_body('6. 客户完整信息查看（含业务联系记录）')
add_body('7. 销售记录添加（自动关联车辆/员工/客户）')
add_body('8. 销售报表生成结果')

add_page_break()


# ============================================================
#  五、用户使用说明
# ============================================================

add_heading_custom('五、用户使用说明', 0)
doc.add_paragraph()

add_heading_custom('5.1 项目编译与运行步骤', 1)
add_body('1. 确保已安装 CMake（3.10+）和 GCC')
add_body('2. 在项目根目录下创建构建目录并编译：')
add_code('''mkdir -p build && cd build
cmake ..
make''')
add_body('3. 运行程序：')
add_code('./auto_sales')
add_body('4. 首次运行时程序会自动创建 data/ 目录用于存储数据文件')

add_heading_custom('5.2 系统各功能使用说明', 1)
add_body('主菜单提供6个选项：')
add_body('选择 1 进入车辆信息管理：支持增删改查和按型号/价格范围筛选')
add_body('选择 2 进入员工信息管理：支持增删改查和按姓名/学历筛选')
add_body('选择 3 进入客户信息管理：支持增删改查、按名称模糊搜索、追加联系记录')
add_body('选择 4 进入销售信息管理：支持添加销售（自动关联）、多条件查询、生成报表')
add_body('选择 5 手动保存所有数据')
add_body('选择 0 退出（自动保存数据）')
doc.add_paragraph()
add_body('各子菜单中，按提示输入对应选项编号即可操作。修改记录时，输入新值会覆盖原值，直接按回车则保留原值。')

add_heading_custom('5.3 软件使用注意事项', 1)
add_body('1. 编号唯一性：添加车辆和员工时需要保证编号不重复，系统会自动检查')
add_body('2. 销售记录关联：添加销售记录时，建议先录入车辆、员工、客户信息，系统会自动关联填充')
add_body('3. 数据保存：退出系统时会自动保存数据，但建议定期手动保存（主菜单选项5）防止意外退出')
add_body('4. 客户编号和销售编号由系统自动生成，无需手动输入')
add_body('5. 业务联系记录追加有长度限制（256字符），超长将被拒绝')
add_body('6. 销售报表导出为 data/report.txt 文本文件')

add_page_break()


# ============================================================
#  六、调试与测试
# ============================================================

add_heading_custom('六、调试与测试', 0)
doc.add_paragraph()

add_heading_custom('6.1 功能测试用例', 1)

# 测试用例表格
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
headers = ['编号', '测试项', '输入', '预期结果', '实际结果']
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for run in p.runs:
            set_run_font(run, '黑体', 'Arial', 10, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

test_cases = [
    ('T1',  '添加车辆',       '编号V001，有效数据',     '添加成功',           '通过'),
    ('T2',  '重复编号',       '编号V001再次添加',       '提示编号已存在',     '通过'),
    ('T3',  '空编号',         '编号输入为空',           '继续等待输入',       '通过'),
    ('T4',  '删除车辆',       '输入V001确认删除',       '删除成功',           '通过'),
    ('T5',  '删除不存在',     '输入不存在的编号',       '提示未找到',         '通过'),
    ('T6',  '修改车辆',       '输入V001，修改价格',     '价格更新成功',       '通过'),
    ('T7',  '回车保留',       '修改时所有字段回车',     '原值保留',           '通过'),
    ('T8',  '按型号查询',     '输入型号关键词',         '显示匹配结果',       '通过'),
    ('T9',  '按价格范围查询', '输入min和max价格',       '显示范围内车辆',     '通过'),
    ('T10', '添加客户',       '输入客户信息',           '自动生成编号C0001',  '通过'),
    ('T11', '追加联系记录',   '输入客户编号和新记录',   '追加成功带日期戳',   '通过'),
    ('T12', '添加销售记录',   '关联已存在的车/员工/客户','自动填充冗余字段',   '通过'),
    ('T13', '销售不关联',     '输入不存在的车辆编号',   '提示手动输入',       '通过'),
    ('T14', '销售报表',       '有多条销售记录',         '统计正确',           '通过'),
    ('T15', '数据持久化',     '添加数据后退出再启动',   '数据完整恢复',       '通过'),
    ('T16', '学历枚举',       '输入1-5以外数字',        '提示重新输入',       '通过'),
    ('T17', '日期范围查询',   '输入起止日期',           '显示范围内记录',     '通过'),
]

for row_data in test_cases:
    row = table.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, '宋体', 'Times New Roman', 10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_heading_custom('6.2 常见bug及解决办法', 1)

bugs = [
    ('1. 缓冲区残留问题',
     '使用 scanf 后直接 fgets 会读到残留换行符。',
     '统一使用 safe_gets() 替代 scanf，内部处理缓冲区清空。'),
    ('2. 删除头节点后 tail 指针未更新',
     '删除唯一节点时 car_tail 仍指向已释放内存。',
     '删除时判断 if (car_tail == p) car_tail = prev。'),
    ('3. 客户编号加载后冲突',
     '重启后 cust_id_counter 为0，新编号与已有编号冲突。',
     '加载时扫描已有编号，更新计数器为最大值。'),
    ('4. 跨平台编译问题',
     'Windows 和 Linux 的 clear 命令和 mkdir 接口不同。',
     '使用 #ifdef _WIN32 条件编译分别处理。'),
]

for title, problem, solution in bugs:
    add_body(title, bold=True, indent=False)
    add_body('问题：' + problem)
    add_body('解决：' + solution)
    doc.add_paragraph()

add_page_break()


# ============================================================
#  七、总结、附录和参考资料
# ============================================================

add_heading_custom('七、总结、附录和参考资料', 0)
doc.add_paragraph()

add_heading_custom('7.1 项目开发总结', 1)
add_body('本次课程设计完成了一个基于C语言的汽车销售信息管理系统。系统涵盖了链表数据结构的核心操作（创建、插入、删除、遍历查找），文件I/O的二进制读写，以及模块化程序设计思想的实际应用。')

add_heading_custom('开发过程中遇到的主要挑战：', 2)
add_body('1. 链表操作的边界条件：头节点删除、尾节点删除、空链表插入等特殊情况需要分别处理，通过维护 head 和 tail 双指针简化了逻辑。')
add_body('2. 跨模块数据关联：销售记录需要引用车辆、员工、客户的编号并自动填充信息，通过暴露各模块的 _find() 函数实现跨模块查询。')
add_body('3. 数据持久化设计：选择二进制文件而非文本文件，因为结构体逐条 fwrite/fread 实现简单且效率较高；但需要在保存时将 next 指针置NULL避免序列化指针地址。')

add_heading_custom('收获与不足：', 2)
add_body('加深了对链表操作的理解，掌握了尾插法、前驱节点追踪、尾指针维护等技巧。')
add_body('学会了模块化设计的接口分离方法（头文件声明＋源文件实现）。')
add_body('不足之处：未实现排序功能、未支持数据导出为可读文本格式（仅报表支持）、无输入数据的合法性深度校验（如价格不能为负）。')

add_heading_custom('7.2 附录', 1)
add_body('（此处附上所有源文件的完整代码，详见仓库 src/ 和 include/ 目录。）')

add_heading_custom('7.3 参考文献', 1)
add_body('[1] 谭浩强. C程序设计（第五版）. 清华大学出版社.')
add_body('[2] K.N. King. C Primer Plus（第六版）. 人民邮电出版社.')
add_body('[3] CMake 官方文档: https://cmake.org/documentation/')
add_body('[4] GNU C Library Manual: https://www.gnu.org/software/libc/manual/')


# ============================================================
#  保存
# ============================================================

output_path = '/home/srk/Git_Program/Automobile_Sales_Information_Management_System/课程设计报告.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
print(f'文件大小：{os.path.getsize(output_path)} bytes')
